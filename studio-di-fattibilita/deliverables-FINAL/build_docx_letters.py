"""
Generate DOCX cover letters per consegna PEC, lettera generica template + 5 varianti
per stakeholder specifici (Coopfond, Regione, ENAC, MIMIT, EASA).

Output:
- 00-LETTERA-Trasmissione-GENERICA.docx
- 01-Lettera-Coopfond.docx
- 02-Lettera-Regione-Liguria.docx
- 03-Lettera-ENAC.docx
- 04-Lettera-MIMIT.docx
- 05-Lettera-EASA.docx
- 06-CdA-1pager.docx
"""

import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

BASE = "/home/user/HALE/studio-di-fattibilita"
ASSETS = "/home/user/HALE/cad"
OUT = os.path.join(BASE, "deliverables-FINAL")
LOGO = os.path.join(ASSETS, "LogoFirmamento Technologies.png")

NAVY = RGBColor(0x0a, 0x1a, 0x3d)
GOLD = RGBColor(0xf0, 0xc9, 0x5c)
GREY = RGBColor(0x3a, 0x3a, 0x3a)

# Dati fissi
FT_NAME = "FIRMAMENTO TECHNOLOGIES SOCIETA' COOPERATIVA"
FT_ADDR = "Via Brigata Liguria 105 R, 16121 Genova (GE)"
FT_FISCAL = "P.IVA / C.F. IT03038500991 . REA 528629"
FT_PEC = "PEC: firmamentotechnologies@pec.it"

def add_logo_header(doc, logo_path=LOGO):
    """Aggiunge header con logo + denominazione."""
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    header = section.header
    htable = header.add_table(rows=1, cols=2, width=Cm(16))
    htable.autofit = False
    htable.columns[0].width = Cm(5)
    htable.columns[1].width = Cm(11)
    # Logo cella sinistra
    cell_logo = htable.cell(0, 0)
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_logo.add_run()
    try:
        run.add_picture(logo_path, width=Cm(4.5))
    except Exception:
        pass
    # Dati cella destra
    cell_data = htable.cell(0, 1)
    cell_data.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = cell_data.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p1.add_run(FT_NAME)
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    p2 = cell_data.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p2.add_run(FT_ADDR)
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    p3 = cell_data.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p3.add_run(f"{FT_FISCAL} . {FT_PEC}")
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    # Linea sotto header
    add_horizontal_line(header.add_paragraph(), NAVY, size=12)

def add_horizontal_line(paragraph, color=NAVY, size=12):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    color_hex = "{:02X}{:02X}{:02X}".format(color[0], color[1], color[2])
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Firmamento Technologies Soc. Coop. . P.IVA IT03038500991 . REA 528629 . PEC firmamentotechnologies@pec.it")
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.color.rgb = GREY

def add_title(doc, text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = NAVY
    return p

def add_para(doc, text, bold=False, italic=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold:
        r.font.bold = True
    if italic:
        r.font.italic = True
    if color:
        r.font.color.rgb = color
    else:
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_destinatario(doc, name, qualifica, addr, pec):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Spett.le")
    r.font.size = Pt(10)
    r.font.italic = True
    for line in [name, qualifica, addr, pec]:
        if line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            r.font.size = Pt(10)
            r.font.bold = (line == name)

def add_oggetto(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    r1 = p.add_run("Oggetto: ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.size = Pt(11)

# ============== TEMPLATE LETTERA BASE ==============
def build_lettera_base(out_path, destinatario, oggetto, corpo_intro, richiesta_specifica=None, allegati_extra=None):
    doc = Document()
    add_logo_header(doc)
    add_footer(doc)

    # Data + luogo
    add_para(doc, "Genova, 17 maggio 2026", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)

    # Destinatario
    dest_filtered = {k: v for k, v in destinatario.items() if k != 'apertura'}
    add_destinatario(doc, **dest_filtered)

    # Oggetto
    add_oggetto(doc, oggetto)

    # Apertura
    add_para(doc, f"Gentile {destinatario.get('apertura', 'Direttore')},", size=11)

    # Corpo
    for paragrafo in corpo_intro:
        add_para(doc, paragrafo, size=11)

    # Struttura Studio
    add_para(doc, "Lo Studio si articola in tre volumi:", size=11)
    p = doc.add_paragraph()
    p.add_run("Volume 1, Studio (testuale): 12 capitoli, dalla Sintesi Esecutiva alla Roadmap post-fattibilita, circa 350 pagine A4.").font.size = Pt(11)
    p = doc.add_paragraph()
    p.add_run("Volume 2, Allegati Tecnici: 13 allegati (RTM, Risk Register, ICD, V&V Plan, Computo Metrico, Safety Case SORA, VIA preliminare e altro), comprensivo di refinement engineering-grade v2.0 (A.4 ICD, A.11 SORA, A.12 VIA) e v1.5 (A.1 RTM, A.2 Risk Register, A.9 Computo, A.10 Manutenzione).").font.size = Pt(11)
    p = doc.add_paragraph()
    p.add_run("Volume 3, Riferimenti: 5 sezioni con circa 270 riferimenti normativi, tecnici e di mercato.").font.size = Pt(11)

    # Caratteristiche
    add_title(doc, "Caratteristiche essenziali del progetto", size=12)
    add_para(doc,
        "Il progetto e sviluppato da Firmamento Technologies in collaborazione con una rete di dieci cooperative aderenti a Legacoop, di cui Fabrica e capofila. L'obiettivo e l'erogazione di servizi ricorrenti, monitoraggio territoriale persistente, connettivita di emergenza, supporto a Protezione Civile e cooperative, rivolti alle Aree Interne italiane. Il focus iniziale e la Regione Liguria, area SNAI Valli Antola-Tigullio, con caso pilota nella frazione di Pentema (Comune di Torriglia, Genova).",
        size=11)
    add_para(doc,
        "La strategia adottata e duale a riduzione del rischio. Il Percorso 6A, VTOL pilota Pentema, copre l'orizzonte 0-12 mesi con budget €700k-€2M baseline e piattaforma commerciale TRL 8-9; il verdetto raccomandato e HOLD CON PIANO REGOLATORIO RAFFORZATO come scenario base (P 45-60%), con possibilita di GO CONDIZIONATO al verificarsi simultaneo di cinque hard conditions. Il Percorso 6B, HALE stratosferico R&D, copre l'orizzonte 24-48+ mesi con budget Phase B €5.5-13.5M R&D-only; il verdetto raccomandato e HOLD CON CRITERI DI USCITA STRINGENTI in attesa dell'apertura del framework EASA HAPS e della disponibilita di partnership con prime contractor.",
        size=11)

    # Richiesta specifica
    if richiesta_specifica:
        add_title(doc, "Richiesta", size=12)
        for r in richiesta_specifica:
            add_para(doc, r, size=11)

    # Chiusura
    add_para(doc, "Restiamo a disposizione per ogni chiarimento, integrazione documentale o approfondimento tecnico.", size=11)
    add_para(doc, "Cordiali saluti.", size=11)

    # Firma
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("Firmamento Technologies Societa Cooperativa")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = NAVY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("[Nome Amministratore Unico]")
    r.font.size = Pt(10)
    r.font.italic = True

    # Allegati
    add_para(doc, "Allegati:", bold=True, size=11)
    allegati_base = [
        "Volume 1 - Studio (PDF firmato digitalmente)",
        "Volume 2 - Allegati Tecnici (PDF firmato digitalmente)",
        "Volume 3 - Riferimenti bibliografici (PDF)",
        "Sintesi Esecutiva 1-pagina per Consiglio di Amministrazione (PDF)",
    ]
    if allegati_extra:
        allegati_base.extend(allegati_extra)
    for a in allegati_base:
        p = doc.add_paragraph(a, style='List Bullet')
        for r in p.runs:
            r.font.size = Pt(10)

    doc.save(out_path)
    print(f"  -> {out_path}")

# ============== DESTINATARI ==============
def main():
    # Lettera generica
    build_lettera_base(
        os.path.join(OUT, "00-LETTERA-Trasmissione-GENERICA.docx"),
        destinatario={
            "name": "[NOME DESTINATARIO]",
            "qualifica": "[Ente / Posizione]",
            "addr": "[Indirizzo]",
            "pec": "[PEC del destinatario]",
            "apertura": "Sig./Direttore/Dott.",
        },
        oggetto="Trasmissione Studio di Fattibilita tecnico-economica, Piattaforma Aerea HALE/VTOL per Aree Interne, Bando Cooding Prototypes Coopfond/Legacoop. Versione M+11 bozza.",
        corpo_intro=[
            "con la presente trasmettiamo, ai fini della Vostra istruttoria, lo Studio di Fattibilita tecnico-economica del progetto \"Piattaforma Aerea HALE/VTOL per Aree Interne\", redatto in conformita all'articolo 41 del D.Lgs. 36/2023 (Codice dei Contratti Pubblici) e al relativo Allegato I.7 (Contenuti minimi del PFTE)."
        ],
        richiesta_specifica=[
            "Prendere visione dello Studio (in particolare Cap. 0 Sintesi Esecutiva, Cap. 10 Raccomandazione di Gate, Cap. 11 Roadmap).",
            "Fornire feedback sulle sezioni di Vostra competenza entro [data da concordare].",
            "Confermare disponibilita a un meeting di approfondimento.",
        ]
    )

    # Lettera Coopfond
    build_lettera_base(
        os.path.join(OUT, "01-Lettera-Coopfond.docx"),
        destinatario={
            "name": "COOPFOND S.p.A.",
            "qualifica": "Direzione Generale, attenzione Responsabile Bando Cooding",
            "addr": "Via Lago di Lugano 19, 00132 Roma",
            "pec": "coopfondspa@pec.legacoop.coop",
            "apertura": "Direttore",
        },
        oggetto="Studio di Fattibilita Piattaforma HALE/VTOL Aree Interne, Bando Cooding Prototypes 2026, Trasmissione versione M+11 bozza.",
        corpo_intro=[
            "con la presente trasmettiamo lo Studio di Fattibilita tecnico-economica del progetto \"Piattaforma Aerea HALE/VTOL per Aree Interne\", presentato in risposta al Bando Cooding Prototypes promosso da Coopfond.",
            "Il progetto si sviluppa attorno a una rete di dieci cooperative aderenti a Legacoop, con Fabrica come capofila, e adotta un modello service-only coerente con la mission cooperativa: nessuna vendita di asset, ma erogazione continuativa di servizi ricorrenti al territorio."
        ],
        richiesta_specifica=[
            "Conferma calendario e dotazione bando Cooding Prototypes 2026 (apertura, scadenze, ammontare massimo per progetto).",
            "Disponibilita a meeting di pre-application entro M+3 per validazione fit del progetto rispetto ai criteri del bando.",
            "Indicazione modalita di accesso a Cooding-Invest per la fase di scale-up Y2-Y3 (€150-300k target).",
            "Verifica disponibilita fondi engagement separati per workshop cooperative pilota (€30-50k stimati)."
        ],
        allegati_extra=[
            "Allegato vendor-rfq: Vendor Quotation Analysis JOUAV vs Tekever",
            "Allegato A.7 financial-model: Modello finanziario Excel"
        ]
    )

    # Lettera Regione Liguria
    build_lettera_base(
        os.path.join(OUT, "02-Lettera-Regione-Liguria.docx"),
        destinatario={
            "name": "REGIONE LIGURIA",
            "qualifica": "Assessorato Innovazione, Ricerca, Universita, attenzione Direzione Generale Sviluppo Economico",
            "addr": "Piazza De Ferrari 1, 16121 Genova",
            "pec": "protocollo@pec.regione.liguria.it",
            "apertura": "Assessore",
        },
        oggetto="Studio di Fattibilita Piattaforma HALE/VTOL per le Aree Interne Liguri, Caso pilota Pentema (Torriglia GE), Richiesta endorsement come anchor customer.",
        corpo_intro=[
            "con la presente trasmettiamo lo Studio di Fattibilita tecnico-economica del progetto \"Piattaforma Aerea HALE/VTOL per Aree Interne\", che individua nella Regione Liguria, e in particolare nell'area SNAI Valli Antola-Tigullio, il territorio pilota di riferimento.",
            "Il caso pilota e la frazione di Pentema (Comune di Torriglia, GE), comunita di quattordici residenti ISTAT in area parco (Parco Naturale Regionale dell'Antola). Il progetto risponde a fabbisogni operativi documentati di Protezione Civile Liguria, ARPA Liguria, Carabinieri Forestali e Comuni SNAI: monitoraggio rischio idrogeologico, antincendio boschivo, connettivita di emergenza, mapping infrastrutture rurali."
        ],
        richiesta_specifica=[
            "Sottoscrizione di Letter of Intent (LoI) o equivalente atto formale entro M+9, come riconoscimento di interesse della Regione al servizio.",
            "Avvio iter per Delibera di Giunta Regionale di endorsement progettuale, eventualmente in collegamento con risorse FESR 2021-2027 (OS 1.1 R&I e OS 5.2 SNAI).",
            "Indicazione referente tecnico per workshop di validazione casi d'uso con Protezione Civile, ARPA, Carabinieri Forestali entro M+3-6.",
            "Disponibilita a meeting di approfondimento con Direzione Generale Sviluppo Economico e Assessorato Innovazione."
        ]
    )

    # Lettera ENAC
    build_lettera_base(
        os.path.join(OUT, "03-Lettera-ENAC.docx"),
        destinatario={
            "name": "ENAC, Ente Nazionale per l'Aviazione Civile",
            "qualifica": "Direzione Regolamentazione Aeroporti e Spazio Aereo, attenzione Ufficio APR",
            "addr": "Viale Castro Pretorio 118, 00185 Roma",
            "pec": "protocollo@pec.enac.gov.it",
            "apertura": "Direttore",
        },
        oggetto="Studio di Fattibilita Piattaforma HALE/VTOL per Aree Interne, Richiesta di pre-application meeting per validazione SAIL preliminare (Specific Category BVLOS), Caso pilota Pentema (Torriglia, GE).",
        corpo_intro=[
            "con la presente trasmettiamo, in via informale e a soli fini istruttori, lo Studio di Fattibilita tecnico-economica del progetto \"Piattaforma Aerea HALE/VTOL per Aree Interne\".",
            "Il Percorso 6A del progetto (pilota VTOL su area Pentema, Comune di Torriglia, GE) prevede operazioni BVLOS in Specific Category con classificazione SAIL preliminare III. La preparazione della SORA application (EASA SORA 2.5, ED Decision 2025/018/R Amendment 3) e in corso, e l'Allegato A.11 dello Studio contiene il Safety Case preliminare completo (Steps 1-8) con i 24 OSO matrice."
        ],
        richiesta_specifica=[
            "Disponibilita a pre-application meeting entro M+3-6 per condivisione del Concept of Operations e validazione preliminare della classificazione SAIL.",
            "Feedback informale sull'approccio tecnico-procedurale, in particolare su: classificazione ground risk per area Pentema (sparse vs moderate), TMPR per Detect-And-Avoid non-cooperative, integrazione spazio aereo Class G mountainous.",
            "Indicazione referente per dialogo continuativo durante la preparazione della SORA application formale."
        ],
        allegati_extra=[
            "Allegato A.11: SORA Safety Case Preliminary v2.0",
            "Allegato A.4: ICD Detailed v2.0 (con dettaglio DAA e Privacy)"
        ]
    )

    # Lettera MIMIT
    build_lettera_base(
        os.path.join(OUT, "04-Lettera-MIMIT.docx"),
        destinatario={
            "name": "MINISTERO DELLE IMPRESE E DEL MADE IN ITALY",
            "qualifica": "Direzione Generale per le politiche industriali, dell'innovazione e delle piccole e medie imprese, Ufficio Aerospazio",
            "addr": "Via Molise 2, 00187 Roma",
            "pec": "dgsii.dg@pec.mise.gov.it",
            "apertura": "Direttore",
        },
        oggetto="Studio di Fattibilita Piattaforma HALE/VTOL per Aree Interne, Posizionamento nella filiera aerospace italiana e nella sovranita tecnologica europea complementare a IRIS².",
        corpo_intro=[
            "con la presente trasmettiamo lo Studio di Fattibilita tecnico-economica del progetto \"Piattaforma Aerea HALE/VTOL per Aree Interne\", proponendo un posizionamento di filiera aerospace italiana coerente con la strategia di sovranita tecnologica europea.",
            "Il progetto si posiziona come potenziale nodo italiano fondatore di una futura infrastruttura sovrana europea HAPS (High Altitude Pseudo-Satellite), complementare a IRIS² sul layer stratosferico. La finestra strategica Y2-Y4 (2027-2030) si sovrappone alla fase pre-operativa di IRIS² (lancio 2029, operativita piena 2031) e apre un'opportunita di posizionamento per un layer stratosferico gap-filler nell'architettura sovrana EU multi-orbit."
        ],
        richiesta_specifica=[
            "Indicazione bandi PNRR Aerospazio e linee IS4Aerospace 2026-2028 applicabili al Percorso 6A (€100-300k target) e al Percorso 6B Phase B R&D (€1-2.5M target).",
            "Endorsement strategico del posizionamento \"complementare a IRIS²\" presso DG CNECT e DG DEFIS della Commissione Europea.",
            "Disponibilita a interlocuzione su programma EU HAPS dedicato (analog IRIS² stratospheric), precondizione per la Fase 5 della visione 10 anni."
        ]
    )

    # Lettera EASA
    build_lettera_base(
        os.path.join(OUT, "05-Lettera-EASA.docx"),
        destinatario={
            "name": "EASA, European Union Aviation Safety Agency",
            "qualifica": "Innovation Network, Drones and U-space Department",
            "addr": "Konrad-Adenauer-Ufer 3, 50668 Cologne, Germany",
            "pec": "innovation@easa.europa.eu",
            "apertura": "Director",
        },
        oggetto="Italian HALE/VTOL Feasibility Study for Inner Areas, Engagement request on EASA HAPS framework and Special Condition dialogue.",
        corpo_intro=[
            "we herewith transmit the Feasibility Study of the project \"HALE/VTOL Aerial Platform for Italian Inner Areas\", developed by Firmamento Technologies in compliance with Italian art. 41 D.Lgs. 36/2023 and NASA Systems Engineering Handbook Rev 2 methodology.",
            "The 6B path of the project (HALE stratospheric solar UAV) faces the well-known regulatory gap of the absence of an EASA HAPS framework. We respectfully request the opportunity to engage with the EASA Innovation Network on the topic of Special Condition HAPS, in line with the ongoing dialogue between the European HAPS industrial ecosystem (CIRA, TAS, Airbus subsidiaries, Sceye) and EASA Rulemaking."
        ],
        richiesta_specifica=[
            "Information on the timeline of any planned RMT (Rulemaking Task) on HAPS perennial operations under EASA Basic Regulation 2018/1139.",
            "Confirmation of availability of Innovation Network office hours for early-stage consultation on the architecture choices of the project (Solar+LiS energy balance, autonomous high-altitude operations, DAA architecture).",
            "Preliminary indication on the feasibility of a Special Condition Light UAS extension or a dedicated Special Condition negotiable for HALE solar configurations."
        ]
    )

    # 06-CdA-1pager.docx
    # Per il 1-pager CdA, conversione da master-deliverables/03-CdA-1pager.md via pandoc DOCX
    import subprocess
    cda_md = os.path.join(BASE, "master-deliverables", "03-CdA-1pager.md")
    cda_docx = os.path.join(OUT, "06-CdA-1pager.docx")
    subprocess.run(
        ["pandoc", "-f", "gfm", "-t", "docx", "-o", cda_docx, cda_md],
        check=True, timeout=30
    )
    print(f"  -> {cda_docx} (via pandoc)")

if __name__ == "__main__":
    main()
